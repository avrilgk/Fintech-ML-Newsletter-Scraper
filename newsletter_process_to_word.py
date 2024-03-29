# -*- coding: utf-8 -*-
"""
Created on Tue Mar 12 14:34:10 2024

@author: SJ
"""

import docx
from docx import Document
import pandas as pd

def replace_text_in_runs(runs, placeholder, replacement):
    for run in runs:
        print(run.text)
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            print("placeholder replaced! "+placeholder)
                    
def replace_placeholder_in_headers_footers(doc, placeholder, replacement):
    for section in doc.sections:
        # Handle default header and footer
        replace_in_header_footer(section.header, placeholder, replacement)
        replace_in_header_footer(section.footer, placeholder, replacement)

        # Handle first page header and footer if they exist
        if section.first_page_header:
            replace_in_header_footer(section.first_page_header, placeholder, replacement)
        if section.first_page_footer:
            replace_in_header_footer(section.first_page_footer, placeholder, replacement)

        # Handle even page header and footer if they exist
        if section.even_page_header:
            replace_in_header_footer(section.even_page_header, placeholder, replacement)
        if section.even_page_footer:
            replace_in_header_footer(section.even_page_footer, placeholder, replacement)

def replace_in_header_footer(header_or_footer, placeholder, replacement):
    for paragraph in header_or_footer.paragraphs:
        replace_text_in_runs(paragraph.runs, placeholder, replacement)


def replace_placeholder_everywhere(doc, placeholder, replacement):
    # Replace in the main document body
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph.runs, placeholder, replacement)
    
    # Replace in headers and footers
    replace_placeholder_in_headers_footers(doc, placeholder, replacement)
    
    
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink



#This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"

def add_articles(doc, articles):
    for paragraph in doc.paragraphs:
        if "start" in paragraph.text:
            # Replace the placeholder with an empty string, effectively removing it
            paragraph.text = paragraph.text.replace("start", "")
            # Add the new text in a new paragraph after the original paragraph
            for article in articles:
                title = doc.add_paragraph(article['title']+" ")
                add_hyperlink(title, f"[{article['source']}]", article['link'])
                
                doc.add_paragraph('- ' + str(article['description']) if article['description'] else '')

def get_records_from_excel(file):
    df = pd.read_excel(file)
    return df.to_dict(orient='records')

def write_newsletter(uploaded_file, pub_date):
    
    doc = Document('Publish Env Template.docx')
    
    articles = get_records_from_excel(uploaded_file)
    
    top_2 = articles[:2]
    
    top_2_headlines = articles[0]['title'] + "; " + articles[1]['title']
    
    remaining = sorted(articles[2:], key=lambda x: x['title'])
    
    ordered_articles = top_2 + remaining
    
    replace_placeholder_everywhere(doc, 'pub', pub_date)
    replace_placeholder_everywhere(doc, 'top_2_headlines', top_2_headlines)
    add_articles(doc, ordered_articles)
    return doc
    
def write_newsletter_from_df(df, pub_date):
    
    doc = Document('Publish Env Template.docx')
    
    articles = df.to_dict(orient='records')
    
    top_2 = articles[:2]
    
    top_2_headlines = articles[0]['title'] + "; " + articles[1]['title']
    
    remaining = sorted(articles[2:], key=lambda x: x['title'])
    
    ordered_articles = top_2 + remaining
    
    replace_placeholder_everywhere(doc, 'pub', pub_date)
    replace_placeholder_everywhere(doc, 'top_2_headlines', top_2_headlines)
    add_articles(doc, ordered_articles)
    return doc
    
     
